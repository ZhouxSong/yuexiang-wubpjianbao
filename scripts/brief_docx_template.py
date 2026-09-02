# -*- coding: utf-8 -*-
"""无 BP 公司简报 Word 参数化模板（公司简报格式规范，蒸馏自真实无BP建档实战脚本）。

用法：
  1. 复制本文件到专题目录，命名 gen_{拼音简称}_report.py；
  2. 仅修改 ===== 数据区 ===== 内的常量与 build() 内的内容调用；
  3. python gen_xxx_report.py 生成 docx。

格式规范（勿改样式函数）：
  Letter 8.5x11in；边距 上下72pt 左右90pt；正文仿宋10.5pt；
  Heading 1/2/3 = 16/14/12pt 加粗黑字；主标题 26pt 加粗深蓝 #1A1A2E 居中；
  表头深蓝底白字加粗单线边框；页眉左公司名右"公司简报"；页脚"第 N 页"(PAGE域)。

已知坑：
  - 中文字体必须同时写 w:rFonts 的 w:eastAsia，仅设 run.font.name 无效；
  - 页码必须用 w:fldChar + w:instrText="PAGE" 域，禁止硬编码数字；
  - 空行用 doc.add_paragraph()，误传 'Normal' 会插入字面量"Normal"；
  - 表格列宽总宽不超过约 15.2cm（Letter - 左右90pt边距）。
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ===== 数据区（每个标的只改这里）=====
COMPANY = '深圳市某某科技有限责任公司'      # 公司全称（页眉左、主标题）
DOC_TITLE = '公司简报'                       # 页眉右、副标题
DATE_LABEL = '2026年9月'                     # 标题块日期
QUERY_DATE = '2026年9月1日'                  # 开篇声明查询日期
OUT_PATH = r'%s公司简报_%s.docx' % (COMPANY, DATE_LABEL)

DARK = RGBColor(0x1A, 0x1A, 0x2E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---------- 页面设置 ----------
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Pt(72)
sec.bottom_margin = Pt(72)
sec.left_margin = Pt(90)
sec.right_margin = Pt(90)


def set_east_asia(owner, font='仿宋'):
    rpr = owner.element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn('w:eastAsia'), font)


# ---------- 样式 ----------
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(10.5)
set_east_asia(normal, '仿宋')

for name, size in (('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 12)):
    st = doc.styles[name]
    st.font.name = '仿宋'
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = BLACK
    set_east_asia(st, '仿宋')
    st.paragraph_format.space_before = Pt(10 if name != 'Heading 1' else 14)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.15


def add_run(p, text, bold=False, size=10.5, color=None, font='仿宋', italic=False):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.name = font
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font)
    if color is not None:
        r.font.color.rgb = color
    return r


def body(text, indent=True, size=10.5):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, size=size)
    return p


def bullet(text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, size=size)
    return p


def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    add_run(p, text, size=9, color=GRAY)
    return p


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def set_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:%s' % edge)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '1A1A2E')
        borders.append(el)
    tblPr.append(borders)


def make_table(headers, rows, font_size=9.5, align_map=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=font_size, color=WHITE)
        shade_cell(c, '1A1A2E')
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            c = t.cell(i, j)
            c.text = ''
            p = c.paragraphs[0]
            al = (align_map or {}).get(j, None)
            if al:
                p.alignment = al
            p.paragraph_format.line_spacing = 1.15
            add_run(p, str(val), size=font_size)
    set_borders(t)
    return t


def page_field(paragraph, size=9):
    run = paragraph.add_run()
    run.font.size = Pt(size)
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(it); run._r.append(f2)


# ---------- 页眉：左公司名 / 右简报标题 ----------
hp = sec.header.paragraphs[0]
hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)
add_run(hp, COMPANY, size=9)
add_run(hp, '\t' + DOC_TITLE, size=9)
pb = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '1A1A2E')
pb.append(bottom)
hp._p.get_or_add_pPr().append(pb)

# ---------- 页脚：居中 第 N 页 ----------
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(fp, '第 ', size=9)
page_field(fp, size=9)
add_run(fp, ' 页', size=9)


def build():
    # ---------- 标题块 ----------
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(4)
    add_run(p, COMPANY, bold=True, size=26, color=DARK)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    add_run(p, DOC_TITLE, size=14, color=DARK, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    add_run(p, DATE_LABEL, size=11)

    # ---------- 开篇声明（强制） ----------
    body('本简报基于公开渠道信息整理（查询日期%s）。需要特别说明：……（按 '
         'references/report-structure.md 第 0 节模板填写：信息透明度描述、渠道列举、'
         '事实与推测分离声明）。' % QUERY_DATE)

    # ---------- 第一章 ----------
    doc.add_heading('第一章 公司概况', level=1)
    doc.add_heading('1.1 工商基本信息', level=2)
    body('（定位描述段）')
    make_table(['项目', '内容'], [
        ['公司名称', COMPANY],
        ['统一社会信用代码', ''],
        ['法定代表人', ''],
        ['成立日期', ''],
        ['注册资本', ''],
        ['企业类型', ''],
        ['登记状态', ''],
        ['注册地址', ''],
        ['所属行业', ''],
        ['人员规模', ''],
        ['官方网站 / 联系方式', ''],
        ['登记机关', ''],
    ], font_size=9.5)
    caption('表1-1 工商基本信息（来源：企查查/天眼查，查询日期%s）' % QUERY_DATE)
    body('经营范围：……（全文）')
    body('经营范围解读：……（业务线索条目+推测标注）')

    doc.add_heading('1.2 股东结构与实际控制关系', level=2)
    make_table(['股东名称', '持股比例', '备注'], [], font_size=9)
    caption('表1-2 股东结构（来源：天眼查股权穿透，查询日期%s）' % QUERY_DATE)
    body('控制权测算：……（穿透计算过程）')

    doc.add_heading('1.3 核心团队与背景线索', level=2)
    make_table(['姓名', '角色', '背景线索'], [], font_size=9)
    caption('表1-3 核心团队线索（均为公开渠道拼合信息，未经公司确认）')

    doc.add_heading('1.4 资质与标签', level=2)
    bullet('（融资阶段/规模/产业归类/资质标签）')

    # ---------- 第二章 ----------
    doc.add_heading('第二章 融资与经营状况', level=1)
    doc.add_heading('2.1 融资情况', level=2)
    make_table(['项目', '内容'], [], font_size=9.5)
    caption('表2-1 融资概况（来源：…，%s）' % QUERY_DATE)
    doc.add_heading('2.2 投资方背景', level=2)
    make_table(['投资方', '背景说明'], [], font_size=9)
    caption('表2-2 投资方背景（来源：…）')
    doc.add_heading('2.3 经营与财务线索', level=2)
    body('（有则写年报/参保/招聘/订单线索；无则逐项"未公开披露"）')

    # ---------- 第三章 ----------
    doc.add_heading('第三章 业务方向与技术逻辑', level=1)
    doc.add_heading('3.1 业务定位', level=2)
    body('（引用权威来源原文+术语科普）')
    doc.add_heading('3.2 技术路线推测（未经公司确认）', level=2)
    bullet('（每条末尾标（推测））')
    body('以上均为行业常识映射下的推测，不构成对公司实际技术路线的认定。')
    doc.add_heading('3.3 知识产权与资产状况', level=2)
    body('（检索结果+截止日期；空白则说明原因并导入尽调）')

    # ---------- 第四章 ----------
    doc.add_heading('第四章 行业与市场分析', level=1)
    doc.add_heading('4.1 市场规模与增长', level=2)
    make_table(['口径', '数据', '来源与时间'], [], font_size=9)
    caption('表4-1 行业关键数据（注：各机构统计口径差异大，引用须以来源为准）')
    doc.add_heading('4.2 竞争格局', level=2)
    make_table(['企业', '总部', '核心方向', '现状'], [], font_size=8.5)
    caption('表4-2 主要玩家（来源：…）')
    doc.add_heading('4.3 政策环境与行业风险信号', level=2)
    bullet('（政策利好+风险信号，逐条注明来源时间）')

    # ---------- 第五章 ----------
    doc.add_heading('第五章 其他重要信息与风险提示', level=1)
    doc.add_heading('5.1 同名/近似名公司辨析（重要）', level=2)
    make_table(['主体', '辨析说明'], [], font_size=9)
    caption('表5-1 同名/近似名主体辨析')
    doc.add_heading('5.2 风险提示', level=2)
    bullet('（极早期/信息不透明/团队核实/投资方身份/行业/注册地）')
    doc.add_heading('5.3 尽调核实建议清单', level=2)
    bullet('（团队/技术/融资/规划/合规五类，每条可执行）')
    doc.add_heading('5.4 主要信息来源清单', level=2)
    make_table(['类别', '来源', '时间', '支撑内容'], [], font_size=8.5)
    caption('表5-2 主要信息来源清单')

    body('')
    body('小结：……（客观中性复述：主体画像→行业位置→信息透明度→尽调重点；不给投资建议）')


build()
doc.save(OUT_PATH)
print('Saved:', OUT_PATH)
